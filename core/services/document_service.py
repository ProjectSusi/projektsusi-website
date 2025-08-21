"""
Document Processing Service
Handles business logic for document management operations
"""

import asyncio
import hashlib
import logging
import mimetypes
import os
import time
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from ..middleware import TenantContext
from ..models.api_models import DocumentResponse, DocumentUpdate
from ..repositories.audit_repository import SwissAuditRepository
from ..repositories.interfaces import IDocumentRepository, IVectorSearchRepository
from ..utils.encryption import get_encryption_manager, is_encryption_enabled

# Import metrics collection
try:
    from ..middleware.metrics_middleware import get_doc_metrics, get_db_metrics
    METRICS_AVAILABLE = True
except ImportError:
    METRICS_AVAILABLE = False

try:
    from ..config.config import config
except ImportError:
    try:
        from config.config import config
    except ImportError:
        config = None

# Import S3 storage service
try:
    from .s3_storage_service import S3StorageService
    S3_AVAILABLE = True
except ImportError:
    S3_AVAILABLE = False

logger = logging.getLogger(__name__)


class DocumentProcessingService:
    """Service for document processing business logic"""

    def __init__(
        self,
        doc_repo: IDocumentRepository,
        vector_repo: IVectorSearchRepository,
        audit_repo: SwissAuditRepository,
        s3_storage: Optional[S3StorageService] = None,
    ):
        self.doc_repo = doc_repo
        self.vector_repo = vector_repo
        self.audit_repo = audit_repo
        self.s3_storage = s3_storage
        
        # Initialize metrics collectors
        if METRICS_AVAILABLE:
            self.doc_metrics = get_doc_metrics()
            self.db_metrics = get_db_metrics()
        else:
            self.doc_metrics = None
            self.db_metrics = None
        
        # Determine storage method
        self.use_s3_storage = (
            self.s3_storage is not None and 
            getattr(config, "USE_S3_STORAGE", False)
        )
        
        if self.use_s3_storage:
            logger.info("Using S3/MinIO storage for documents")
        else:
            logger.info("Using local filesystem storage for documents")

        # File validation settings
        self.max_file_size = getattr(config, "MAX_FILE_SIZE", 50 * 1024 * 1024)  # 50MB
        self.allowed_extensions = {".pdf", ".docx", ".txt", ".md", ".csv", ".xlsx"}
        self.allowed_mime_types = {
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "text/plain",
            "text/markdown",
            "text/csv",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        }

        # Content filtering settings
        self.problematic_keywords = [
            "zero-hallucination",
            "guidelines for following",
            "only use information",
            "additional guidelines",
            "training instructions",
            "quelels",
        ]
        self.bio_waste_keywords = [
            "bioabfall",
            "bio waste",
            "organic waste",
            "kompost",
            "grünabfall",
            "küchenabfälle",
            "obst",
            "gemüse",
            "fruit",
            "vegetable",
            "food waste",
        ]

    async def _store_document_content(
        self, content: bytes, filename: str, tenant_id: int
    ) -> Tuple[str, Optional[bytes]]:
        """Store document content either locally or in S3, with optional encryption"""
        
        # Generate unique filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        stored_filename = f"{timestamp}_{filename}"
        
        # Handle encryption if enabled
        encryption_salt = None
        content_to_store = content
        
        if is_encryption_enabled():
            try:
                encryption_manager = get_encryption_manager()
                encrypted_content, encryption_salt = (
                    encryption_manager.encrypt_document_content(content, tenant_id)
                )
                content_to_store = encrypted_content
                logger.info(f"Document encrypted for tenant {tenant_id}")
            except Exception as e:
                logger.error(f"Encryption failed, storing unencrypted: {e}")
                content_to_store = content
        
        if self.use_s3_storage:
            # Store in S3/MinIO
            try:
                # Generate S3 object key
                object_key = self.s3_storage._generate_object_key(
                    tenant_id=tenant_id,
                    document_id=0,  # Will be updated after document creation
                    filename=stored_filename,
                    document_type="upload"
                )
                
                # Upload to S3
                metadata = {
                    "tenant_id": str(tenant_id),
                    "original_filename": filename,
                    "encrypted": str(is_encryption_enabled()),
                    "upload_timestamp": datetime.now().isoformat()
                }
                
                upload_result = await self.s3_storage.upload_document(
                    content_to_store,
                    object_key,
                    content_type="application/octet-stream",
                    metadata=metadata
                )
                
                # Return S3 object key as file_path for storage in database
                file_path = f"s3://{self.s3_storage.bucket_name}/{object_key}"
                logger.info(f"Document uploaded to S3: {file_path}")
                
                return file_path, encryption_salt
                
            except Exception as e:
                logger.error(f"S3 upload failed, falling back to local storage: {e}")
                # Fall through to local storage
        
        # Store locally
        upload_dir = Path(
            config.UPLOAD_DIR
            if config and hasattr(config, "UPLOAD_DIR")
            else "data/storage/uploads"
        )
        upload_dir.mkdir(parents=True, exist_ok=True)
        
        file_path = upload_dir / stored_filename
        
        # Write to local file
        with open(file_path, "wb") as f:
            f.write(content_to_store)
        
        logger.info(f"Document stored locally: {file_path}")
        return str(file_path), encryption_salt

    async def validate_upload(
        self, filename: str, content: bytes, content_type: str
    ) -> Tuple[bool, str]:
        """Validate uploaded file"""
        try:
            # Check filename
            if not filename or filename.strip() == "":
                return False, "Filename cannot be empty"

            # Check file extension
            file_ext = Path(filename).suffix.lower()
            if file_ext not in self.allowed_extensions:
                return (
                    False,
                    f"File type {file_ext} not allowed. Allowed: {', '.join(self.allowed_extensions)}",
                )

            # Check file size
            if len(content) > self.max_file_size:
                return (
                    False,
                    f"File size {len(content)} exceeds maximum {self.max_file_size} bytes",
                )

            # Check MIME type
            detected_type = mimetypes.guess_type(filename)[0]
            if detected_type not in self.allowed_mime_types:
                return False, f"MIME type {detected_type} not allowed"

            # Basic content validation
            if len(content) == 0:
                return False, "File cannot be empty"

            # Security: Check for path traversal and suspicious filenames
            if any(sequence in filename for sequence in ["..", "..\\", "../", "..\\"]):
                return False, "Path traversal attempt detected in filename"

            if any(char in filename for char in ["\x00", "\r", "\n", "\t"]):
                return False, "Control characters detected in filename"

            # Additional checks for Windows/Unix reserved names
            reserved_names = {
                "CON",
                "PRN",
                "AUX",
                "NUL",
                "COM1",
                "COM2",
                "COM3",
                "COM4",
                "COM5",
                "COM6",
                "COM7",
                "COM8",
                "COM9",
                "LPT1",
                "LPT2",
                "LPT3",
                "LPT4",
                "LPT5",
                "LPT6",
                "LPT7",
                "LPT8",
                "LPT9",
            }
            base_name = Path(filename).stem.upper()
            if base_name in reserved_names:
                return False, f"Reserved filename '{base_name}' not allowed"

            return True, "Valid file"

        except Exception as e:
            logger.error(f"File validation error: {e}")
            return False, f"Validation failed: {str(e)}"

    def _generate_file_hash(self, content: bytes) -> str:
        """Generate SHA-256 hash of file content"""
        return hashlib.sha256(content).hexdigest()

    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename for safe storage - SECURE against path traversal"""
        # Remove path components completely (security critical)
        filename = os.path.basename(filename)

        # Remove any remaining path separators and dangerous sequences
        filename = filename.replace("..", "").replace("/", "").replace("\\", "")
        filename = filename.replace("\x00", "")  # Null byte removal

        # Replace problematic characters
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            filename = filename.replace(char, "_")

        # Remove control characters
        filename = "".join(char for char in filename if ord(char) >= 32)

        # Ensure filename doesn't start with dot (hidden files)
        if filename.startswith("."):
            filename = "_" + filename[1:]

        # Ensure it's not empty after sanitization
        if not filename or filename.strip() == "":
            filename = f"document_{int(datetime.now().timestamp())}"

        # Limit length
        if len(filename) > 255:
            name, ext = os.path.splitext(filename)
            filename = name[: 250 - len(ext)] + ext

        return filename

    async def process_upload(
        self,
        filename: str,
        content: bytes,
        content_type: str,
        uploader_id: Optional[str] = None,
    ) -> DocumentResponse:
        """Process document upload with full business logic"""
        try:
            # Validate upload
            is_valid, validation_message = await self.validate_upload(
                filename, content, content_type
            )
            if not is_valid:
                # Log failed validation attempt
                from ..repositories.audit_repository import AuditEntry, AuditEventType, DataClassification
                audit_entry = AuditEntry(
                    event_type=AuditEventType.DOCUMENT_UPLOAD,
                    user_id=uploader_id,
                    action_description=f"Document upload validation failed: {validation_message}",
                    resource_accessed=filename,
                    response_status=400,
                    data_classification=DataClassification.INTERNAL,
                    metadata={"validation_error": validation_message}
                )
                await self.audit_repo.log_event(audit_entry)
                raise ValueError(validation_message)

            # Sanitize filename
            safe_filename = self._sanitize_filename(filename)

            # Generate file hash for deduplication
            file_hash = self._generate_file_hash(content)

            # Check for duplicates
            existing = await self.doc_repo.find_by_hash(file_hash)
            if existing:
                logger.info(f"Duplicate file detected: {safe_filename}")
                # Log duplicate detection
                audit_entry = AuditEntry(
                    event_type=AuditEventType.DOCUMENT_UPLOAD,
                    user_id=uploader_id,
                    document_id=existing.id,
                    action_description=f"Duplicate document upload detected: {safe_filename}",
                    resource_accessed=filename,
                    response_status=200,
                    data_classification=DataClassification.INTERNAL,
                    metadata={"duplicate_of": existing.id, "file_hash": file_hash}
                )
                await self.audit_repo.log_event(audit_entry)
                return DocumentResponse(
                    id=existing.id,
                    filename=existing.filename,
                    size=existing.file_size,
                    content_type=existing.content_type,
                    status="duplicate",
                    message="File already exists",
                )

            # Save file to storage (local or S3)
            tenant_id = TenantContext.get_current_tenant_id()
            file_path, encryption_salt = await self._store_document_content(
                content, safe_filename, tenant_id
            )

            # Create document record
            from ..repositories.models import Document, DocumentStatus

            # Prepare metadata
            metadata = {"file_hash": file_hash, "encrypted": is_encryption_enabled()}

            # Add encryption salt to metadata if encryption is used
            if encryption_salt:
                import base64

                metadata["encryption_salt"] = base64.b64encode(encryption_salt).decode()

            document = Document(
                tenant_id=tenant_id,
                filename=safe_filename,
                original_filename=filename,
                file_path=str(file_path),
                file_size=len(content),
                content_type=content_type,
                uploader=uploader_id or "anonymous",
                upload_timestamp=datetime.now(),
                status=DocumentStatus.UPLOADING,
                metadata=metadata,
            )

            # Store in repository
            document = await self.doc_repo.create(document)

            # Start async processing
            asyncio.create_task(self._process_document_async(document.id, file_path))

            # Log successful upload
            audit_entry = AuditEntry(
                event_type=AuditEventType.DOCUMENT_UPLOAD,
                user_id=uploader_id,
                document_id=document.id,
                action_description=f"Document uploaded successfully: {safe_filename}",
                resource_accessed=filename,
                response_status=200,
                data_classification=DataClassification.INTERNAL,
                metadata={
                    "file_size": len(content),
                    "content_type": content_type,
                    "encrypted": is_encryption_enabled(),
                    "tenant_id": tenant_id
                }
            )
            await self.audit_repo.log_event(audit_entry)
            
            logger.info(
                f"Document uploaded successfully: {safe_filename} (ID: {document.id})"
            )

            return DocumentResponse(
                id=document.id,
                filename=document.filename,
                size=document.file_size,
                content_type=document.content_type,
                status="uploaded",
                message="Document uploaded successfully",
            )

        except Exception as e:
            logger.error(f"Document upload processing failed: {e}")
            # Log upload failure
            audit_entry = AuditEntry(
                event_type=AuditEventType.DOCUMENT_UPLOAD,
                user_id=uploader_id,
                action_description=f"Document upload failed: {str(e)}",
                resource_accessed=filename,
                response_status=500,
                data_classification=DataClassification.INTERNAL,
                metadata={"error": str(e)}
            )
            await self.audit_repo.log_event(audit_entry)
            raise

    async def _process_document_async(self, document_id: int, file_path: Path):
        """Process document in background (chunking, embedding, etc.)"""
        try:
            # Update status to processing
            await self.doc_repo.update_status(document_id, "processing")

            # Extract text from document
            text_content = await self._extract_text(file_path)
            if not text_content:
                raise ValueError("No text content extracted from document")

            # Update document with text content
            await self.doc_repo.update(document_id, {"text_content": text_content})

            # Create chunks
            chunks = self._create_chunks(text_content, chunk_size=500, overlap=50)
            logger.info(f"Created {len(chunks)} chunks for document {document_id}")

            # Store chunks in database
            chunk_ids = await self._store_chunks(document_id, chunks)

            # Generate embeddings for chunks
            embeddings = await self._generate_embeddings(chunks)

            # Store embeddings in database
            embedding_records = await self._store_embeddings(
                document_id, chunk_ids, embeddings
            )

            # Update vector index
            await self.vector_repo.add_to_index(embedding_records)

            # Update document status and counts
            await self.doc_repo.update(
                document_id,
                {
                    "status": "completed",
                    "chunk_count": len(chunks),
                    "embedding_count": len(embeddings),
                    "completion_timestamp": datetime.now(),
                },
            )

            logger.info(
                f"Document {document_id} processed successfully with {len(chunks)} chunks"
            )

        except Exception as e:
            logger.error(f"Document processing failed for {document_id}: {e}")
            await self.doc_repo.update_status(document_id, "failed")

    async def get_document_details(self, document_id: int) -> Dict[str, Any]:
        """Get detailed document information"""
        try:
            document = await self.doc_repo.get_by_id(document_id)
            if not document:
                raise ValueError(f"Document {document_id} not found")

            return {
                "document": {
                    "id": document.id,
                    "filename": document.filename,
                    "original_filename": getattr(
                        document, "original_filename", document.filename
                    ),
                    "size": document.file_size,
                    "content_type": document.content_type,
                    "status": document.status,
                    "upload_date": (
                        document.upload_date.isoformat()
                        if document.upload_date
                        else None
                    ),
                    "uploader": getattr(document, "uploader", "unknown"),
                    "hash": getattr(document, "file_hash", None),
                }
            }

        except Exception as e:
            logger.error(f"Error getting document details: {e}")
            raise

    async def update_document(
        self, document_id: int, updates: DocumentUpdate
    ) -> Dict[str, Any]:
        """Update document metadata"""
        try:
            # Get existing document
            document = await self.doc_repo.get_by_id(document_id)
            if not document:
                raise ValueError(f"Document {document_id} not found")

            # Prepare update data
            update_data = updates.dict(exclude_unset=True)

            # Apply updates
            updated_document = await self.doc_repo.update(document_id, update_data)

            # Note: Audit logging would go here in production
            logger.info(f"Document {document_id} updated successfully")

            return {
                "message": "Document updated successfully",
                "document": {"id": updated_document.id, "updates": update_data},
            }

        except Exception as e:
            logger.error(f"Error updating document: {e}")
            raise

    async def delete_document(self, document_id: int) -> Dict[str, Any]:
        """Delete document and cleanup"""
        try:
            # Get document for cleanup
            document = await self.doc_repo.get_by_id(document_id)
            if not document:
                raise ValueError(f"Document {document_id} not found")

            # Delete from vector store if exists
            try:
                await self.vector_repo.delete_document(document_id)
            except Exception as e:
                logger.warning(f"Vector cleanup failed for document {document_id}: {e}")

            # Delete physical file
            if hasattr(document, "file_path") and document.file_path:
                try:
                    file_path = Path(document.file_path)
                    if file_path.exists():
                        file_path.unlink()
                        logger.info(f"Deleted file: {file_path}")
                except Exception as e:
                    logger.warning(f"File cleanup failed: {e}")

            # Delete from repository
            await self.doc_repo.delete(document_id)

            # Log successful deletion
            audit_entry = AuditEntry(
                event_type=AuditEventType.DOCUMENT_DELETE,
                document_id=document_id,
                action_description=f"Document deleted successfully: {document.filename}",
                resource_accessed=document.filename,
                response_status=200,
                data_classification=DataClassification.INTERNAL,
                metadata={
                    "file_path": document.file_path,
                    "tenant_id": getattr(document, 'tenant_id', 1)
                }
            )
            await self.audit_repo.log_event(audit_entry)
            
            logger.info(f"Document {document_id} deleted successfully")

            return {"message": f"Document {document_id} deleted successfully"}

        except Exception as e:
            logger.error(f"Error deleting document: {e}")
            raise

    async def get_download_path(self, document_id: int) -> str:
        """Get file path for document download - handles both S3 and local storage"""
        try:
            document = await self.doc_repo.get_by_id(document_id)
            if not document:
                raise ValueError(f"Document {document_id} not found")

            if not hasattr(document, "file_path") or not document.file_path:
                raise ValueError(f"No file path for document {document_id}")
            
            # Check if document is stored in S3
            if document.file_path.startswith('s3://'):
                return await self._handle_s3_download(document)
            else:
                return await self._handle_local_download(document)
        
        except Exception as e:
            logger.error(f"Error getting download path: {e}")
            raise
    
    async def _handle_s3_download(self, document) -> str:
        """Handle S3 document download with decryption support"""
        try:
            if not self.s3_storage:
                raise ValueError("S3 storage not configured")
            
            # Extract object key from S3 path
            s3_path = document.file_path  # e.g., "s3://bucket/tenant_1/uploads/file.txt"
            object_key = s3_path.split('/', 3)[-1]  # Extract key after bucket name
            
            # Download content from S3
            content = await self.s3_storage.get_document(object_key)
            
            # Handle decryption if needed
            if document.metadata.get("encrypted", False) and is_encryption_enabled():
                try:
                    salt_b64 = document.metadata.get("encryption_salt")
                    if salt_b64:
                        import base64
                        salt = base64.b64decode(salt_b64.encode())
                        tenant_id = getattr(document, 'tenant_id', 1)
                        
                        encryption_manager = get_encryption_manager()
                        content = encryption_manager.decrypt_document_content(
                            content, salt, tenant_id
                        )
                        
                        # Log successful S3 download with decryption
                        audit_entry = AuditEntry(
                            event_type=AuditEventType.DOCUMENT_DOWNLOAD,
                            document_id=document.id,
                            action_description=f"S3 document decrypted and downloaded: {document.filename}",
                            resource_accessed=document.filename,
                            response_status=200,
                            data_classification=DataClassification.INTERNAL,
                            metadata={"storage": "s3", "encrypted": True, "tenant_id": tenant_id}
                        )
                        await self.audit_repo.log_event(audit_entry)
                        
                except Exception as e:
                    logger.error(f"S3 decryption failed: {e}")
                    # Continue with encrypted content
            else:
                # Log unencrypted S3 download
                audit_entry = AuditEntry(
                    event_type=AuditEventType.DOCUMENT_DOWNLOAD,
                    document_id=document.id,
                    action_description=f"S3 document downloaded: {document.filename}",
                    resource_accessed=document.filename,
                    response_status=200,
                    data_classification=DataClassification.INTERNAL,
                    metadata={"storage": "s3", "encrypted": False}
                )
                await self.audit_repo.log_event(audit_entry)
            
            # Create temporary file for download
            import tempfile
            temp_file = tempfile.NamedTemporaryFile(delete=False)
            temp_file.write(content)
            temp_file.close()
            
            logger.info(f"Downloaded S3 document {document.id} to temporary file")
            return temp_file.name
            
        except Exception as e:
            logger.error(f"S3 download failed: {e}")
            raise
    
    async def _handle_local_download(self, document) -> str:
        """Handle local document download with security validation"""
        try:
            # SECURITY: Resolve path and validate it's within allowed directories
            file_path = Path(document.file_path).resolve()

            # Get allowed storage directories
            upload_dir = Path(
                config.UPLOAD_DIR
                if config and hasattr(config, "UPLOAD_DIR")
                else "data/storage/uploads"
            ).resolve()
            processed_dir = Path(
                config.PROCESSED_DIR
                if config and hasattr(config, "PROCESSED_DIR")
                else "data/storage/processed"
            ).resolve()

            # Check if file is within allowed directories (prevent path traversal)
            is_in_upload = upload_dir in file_path.parents or file_path == upload_dir
            is_in_processed = (
                processed_dir in file_path.parents or file_path == processed_dir
            )

            if not (is_in_upload or is_in_processed):
                logger.error(f"Path traversal attempt detected: {file_path}")
                raise ValueError("Access denied: File path outside allowed directories")

            if not file_path.exists():
                raise ValueError(f"File not found: {file_path}")

            # Additional security: ensure it's a file, not a directory or symlink
            if not file_path.is_file():
                raise ValueError(f"Path is not a regular file: {file_path}")

            # Handle decryption if document is encrypted
            if document.metadata.get("encrypted", False) and is_encryption_enabled():
                try:
                    # Get encryption salt from metadata
                    salt_b64 = document.metadata.get("encryption_salt")
                    if not salt_b64:
                        logger.warning(
                            f"Document {document_id} marked as encrypted but no salt found"
                        )
                        return file_path

                    import base64

                    salt = base64.b64decode(salt_b64.encode())
                    tenant_id = document.tenant_id

                    # Read encrypted file
                    with open(file_path, "rb") as f:
                        encrypted_content = f.read()

                    # Decrypt content
                    encryption_manager = get_encryption_manager()
                    decrypted_content = encryption_manager.decrypt_document_content(
                        encrypted_content, salt, tenant_id
                    )

                    # Create temporary decrypted file for download
                    temp_path = file_path.with_suffix(".tmp")
                    with open(temp_path, "wb") as f:
                        f.write(decrypted_content)

                    # Log successful decryption for download
                    audit_entry = AuditEntry(
                        event_type=AuditEventType.DOCUMENT_DOWNLOAD,
                        document_id=document_id,
                        action_description=f"Document decrypted and downloaded: {document.filename}",
                        resource_accessed=document.filename,
                        response_status=200,
                        data_classification=DataClassification.INTERNAL,
                        metadata={
                            "encrypted": True,
                            "tenant_id": document.tenant_id
                        }
                    )
                    await self.audit_repo.log_event(audit_entry)
                    
                    logger.info(f"Decrypted document {document_id} for download")
                    return temp_path

                except Exception as e:
                    logger.error(f"Failed to decrypt document {document_id}: {e}")
                    # Log decryption failure
                    audit_entry = AuditEntry(
                        event_type=AuditEventType.DOCUMENT_DOWNLOAD,
                        document_id=document_id,
                        action_description=f"Document decryption failed: {str(e)}",
                        resource_accessed=document.filename,
                        response_status=500,
                        data_classification=DataClassification.INTERNAL,
                        metadata={"error": str(e), "tenant_id": document.tenant_id}
                    )
                    await self.audit_repo.log_event(audit_entry)
                    # Return encrypted file as fallback
                    return file_path
            else:
                # Log unencrypted download
                audit_entry = AuditEntry(
                    event_type=AuditEventType.DOCUMENT_DOWNLOAD,
                    document_id=document_id,
                    action_description=f"Document downloaded: {document.filename}",
                    resource_accessed=document.filename,
                    response_status=200,
                    data_classification=DataClassification.INTERNAL,
                    metadata={
                        "encrypted": False,
                        "tenant_id": getattr(document, 'tenant_id', 1)
                    }
                )
                await self.audit_repo.log_event(audit_entry)

            return file_path

        except Exception as e:
            logger.error(f"Error getting download path: {e}")
            raise

    async def _extract_text(self, file_path: Path) -> str:
        """Extract text from document based on file type"""
        try:
            # Check if this is an encrypted file by looking for .enc extension
            # or checking document metadata for encryption flag
            decrypted_content = None
            is_encrypted_file = False
            
            # First, try to determine if file is encrypted
            if file_path.suffix == '.enc' or self._is_file_encrypted(file_path):
                is_encrypted_file = True
                logger.debug(f"Detected encrypted file: {file_path}")
                
                if is_encryption_enabled():
                    try:
                        # Read encrypted file and decrypt content
                        with open(file_path, "rb") as f:
                            encrypted_content = f.read()
                        
                        # Get document metadata to find encryption salt
                        document_id = self._extract_document_id_from_path(file_path)
                        if document_id:
                            document = await self.doc_repo.get_by_id(document_id)
                            if document and document.metadata.get("encrypted", False):
                                salt_b64 = document.metadata.get("encryption_salt")
                                if salt_b64:
                                    import base64
                                    salt = base64.b64decode(salt_b64.encode())
                                    tenant_id = document.tenant_id
                                    
                                    encryption_manager = get_encryption_manager()
                                    decrypted_content = encryption_manager.decrypt_document_content(
                                        encrypted_content, salt, tenant_id
                                    )
                                    logger.debug(f"Successfully decrypted file: {file_path}")
                        
                        if decrypted_content is None:
                            logger.warning(f"Could not decrypt file {file_path}, proceeding with encrypted content")
                            
                    except Exception as e:
                        logger.error(f"Failed to decrypt file {file_path}: {e}")
                        # Continue with original file content
                else:
                    logger.warning(f"File appears encrypted but encryption is disabled: {file_path}")
            
            # Get actual file extension (remove .enc if present)
            actual_file_path = file_path
            file_ext = file_path.suffix.lower()
            if file_ext == '.enc':
                # Get the real extension from the filename before .enc
                base_name = file_path.stem
                if '.' in base_name:
                    file_ext = '.' + base_name.split('.')[-1].lower()
                else:
                    file_ext = '.txt'  # Default fallback

            if file_ext == ".txt" or file_ext == ".md":
                # Read plain text files
                if decrypted_content is not None:
                    return decrypted_content.decode('utf-8', errors='ignore')
                else:
                    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                        return f.read()

            elif file_ext == ".pdf":
                # Extract from PDF
                try:
                    import PyPDF2
                    from io import BytesIO

                    text = ""
                    if decrypted_content is not None:
                        # Use decrypted content
                        pdf_reader = PyPDF2.PdfReader(BytesIO(decrypted_content))
                    else:
                        # Read file directly
                        with open(file_path, "rb") as f:
                            pdf_reader = PyPDF2.PdfReader(f)
                    
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                    return text.strip()
                except ImportError:
                    logger.warning("PyPDF2 not installed, trying fallback")
                    return f"PDF Document: {file_path.name}"

            elif file_ext == ".docx":
                # Extract from Word documents
                try:
                    from docx import Document as DocxDocument
                    from io import BytesIO

                    if decrypted_content is not None:
                        # Use decrypted content
                        doc = DocxDocument(BytesIO(decrypted_content))
                    else:
                        # Read file directly
                        doc = DocxDocument(file_path)
                    
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    return text.strip()
                except ImportError:
                    logger.warning("python-docx not installed, trying fallback")
                    return f"Word Document: {file_path.name}"

            elif file_ext in [".csv", ".xlsx"]:
                # Extract from spreadsheets
                try:
                    import pandas as pd
                    from io import BytesIO

                    if decrypted_content is not None:
                        # Use decrypted content
                        if file_ext == ".csv":
                            df = pd.read_csv(BytesIO(decrypted_content))
                        else:
                            df = pd.read_excel(BytesIO(decrypted_content))
                    else:
                        # Read file directly
                        if file_ext == ".csv":
                            df = pd.read_csv(file_path)
                        else:
                            df = pd.read_excel(file_path)
                    
                    # Convert dataframe to text
                    return df.to_string()
                except ImportError:
                    logger.warning("pandas not installed, trying fallback")
                    return f"Spreadsheet: {file_path.name}"

            else:
                # Unknown file type - try to read as text
                if decrypted_content is not None:
                    return decrypted_content.decode('utf-8', errors='ignore')
                else:
                    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                        return f.read()

        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            # Return at least the filename so we have something
            return f"Document: {file_path.name}"

    def _is_file_encrypted(self, file_path: Path) -> bool:
        """Check if a file appears to be encrypted"""
        try:
            # Simple heuristic: encrypted files typically have high entropy
            # and won't contain readable text patterns
            with open(file_path, "rb") as f:
                sample = f.read(1024)  # Read first 1KB
            
            # If file is very small, assume not encrypted
            if len(sample) < 100:
                return False
            
            # Check for high ratio of non-printable characters
            printable_chars = sum(1 for byte in sample if 32 <= byte <= 126)
            ratio = printable_chars / len(sample)
            
            # If less than 50% printable characters, likely encrypted
            return ratio < 0.5
            
        except Exception:
            return False

    def _extract_document_id_from_path(self, file_path: Path) -> Optional[int]:
        """Extract document ID from file path for metadata lookup"""
        try:
            # This is a simple approach - in production you might want a more robust mapping
            # For now, we'll need to search by file path in the database
            return None  # Will be implemented when needed
        except Exception:
            return None

    def _create_chunks(
        self, text: str, chunk_size: int = 500, overlap: int = 50
    ) -> List[str]:
        """Split text into overlapping chunks"""
        if not text:
            return []

        chunks = []
        words = text.split()

        if len(words) <= chunk_size:
            # Text is small enough to be a single chunk
            return [text]

        # Create overlapping chunks
        for i in range(0, len(words), chunk_size - overlap):
            chunk_words = words[i : i + chunk_size]
            chunk_text = " ".join(chunk_words)
            if chunk_text.strip():  # Only add non-empty chunks
                chunks.append(chunk_text)

        return chunks

    async def _store_chunks(self, document_id: int, chunks: List[str]) -> List[int]:
        """Store text chunks in database"""
        chunk_ids = []

        # Get chunk repository (we need to add this to interfaces)
        # For now, store in SQLite directly
        import sqlite3

        db_path = getattr(config, "DATABASE_PATH", None) if config else None
        conn = sqlite3.connect(db_path or "data/rag_database.db")

        try:
            for i, chunk_text in enumerate(chunks):
                cursor = conn.execute(
                    """
                    INSERT INTO chunks (document_id, chunk_index, text,
                                      character_count, word_count)
                    VALUES (?, ?, ?, ?, ?)
                """,
                    (
                        document_id,
                        i,
                        chunk_text,
                        len(chunk_text),
                        len(chunk_text.split()),
                    ),
                )
                chunk_ids.append(cursor.lastrowid)

            conn.commit()
            return chunk_ids

        finally:
            conn.close()

    async def _generate_embeddings(self, chunks: List[str]) -> List[List[float]]:
        """Generate embeddings for text chunks"""
        try:
            # Use sentence transformers
            from sentence_transformers import SentenceTransformer

            # Initialize model (this should be cached in production)
            model = SentenceTransformer("all-MiniLM-L6-v2")

            # Generate embeddings
            embeddings = model.encode(chunks, convert_to_numpy=True)

            return embeddings.tolist()

        except ImportError:
            logger.error("sentence-transformers not installed")
            # Return dummy embeddings (not for security purposes)
            import secrets

            return [
                [secrets.SystemRandom().random() for _ in range(384)] for _ in chunks
            ]

    async def _store_embeddings(
        self, document_id: int, chunk_ids: List[int], embeddings: List[List[float]]
    ) -> List:
        """Store embeddings in database"""
        from ..repositories.models import Embedding

        embedding_records = []

        # Store in SQLite directly for now
        import gzip
        import pickle
        import sqlite3

        db_path = getattr(config, "DATABASE_PATH", None) if config else None
        conn = sqlite3.connect(db_path or "data/rag_database.db")

        try:
            for chunk_id, embedding_vector in zip(chunk_ids, embeddings):
                # Compress embedding vector
                compressed = gzip.compress(pickle.dumps(embedding_vector))

                cursor = conn.execute(
                    """
                    INSERT INTO embeddings (chunk_id, embedding_data,
                                          embedding_model, dimensions)
                    VALUES (?, ?, ?, ?)
                """,
                    (chunk_id, compressed, "all-MiniLM-L6-v2", len(embedding_vector)),
                )

                # Create Embedding object for vector index
                embedding_obj = Embedding(
                    id=cursor.lastrowid,
                    chunk_id=chunk_id,
                    document_id=document_id,
                    embedding_vector=embedding_vector,
                    embedding_model="all-MiniLM-L6-v2",
                    vector_dimension=len(embedding_vector),
                )
                embedding_records.append(embedding_obj)

            conn.commit()
            return embedding_records

        finally:
            conn.close()

    def analyze_document_content(self, text_content: str) -> Dict[str, Any]:
        """Analyze document content for problematic patterns"""
        content_lower = text_content.lower()

        # Check for problematic content
        problematic_score = sum(
            1 for keyword in self.problematic_keywords if keyword in content_lower
        )

        # Check for bio waste content
        bio_waste_score = sum(
            1 for keyword in self.bio_waste_keywords if keyword in content_lower
        )

        # Check for encoding issues
        corruption_score = text_content.count("�")

        # Classify content type
        content_type = "unknown"
        if bio_waste_score >= 2:
            content_type = "bio_waste"
        elif problematic_score > 0:
            content_type = "training_instructions"
        elif any(
            cs in content_lower for cs in ["javascript", "programming", "software"]
        ):
            content_type = "computer_science"

        is_problematic = (
            problematic_score > 0
            or (corruption_score > 10)
            or (len(text_content.strip()) < 100 and content_type != "bio_waste")
        )

        return {
            "content_type": content_type,
            "is_problematic": is_problematic,
            "bio_waste_score": bio_waste_score,
            "problematic_score": problematic_score,
            "corruption_score": corruption_score,
            "content_length": len(text_content),
            "recommendation": "reject" if is_problematic else "accept",
        }

    async def validate_document_content(
        self, text_content: str
    ) -> Tuple[bool, str, Dict[str, Any]]:
        """Validate document content and return analysis"""
        analysis = self.analyze_document_content(text_content)

        if analysis["is_problematic"]:
            reasons = []
            if analysis["problematic_score"] > 0:
                reasons.append("contains training instructions")
            if analysis["corruption_score"] > 10:
                reasons.append("has encoding corruption")
            if analysis["content_length"] < 100:
                reasons.append("content too short")

            return False, f"Document rejected: {', '.join(reasons)}", analysis

        if analysis["content_type"] == "bio_waste":
            return True, "Bio waste document accepted", analysis
        elif analysis["content_type"] == "unknown" and len(text_content.strip()) > 200:
            return True, "General document accepted", analysis
        else:
            return (
                False,
                "Document type not suitable for bio waste RAG system",
                analysis,
            )
